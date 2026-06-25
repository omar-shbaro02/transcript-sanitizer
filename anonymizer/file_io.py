from pathlib import Path
from typing import Iterable

SUPPORTED_EXTENSIONS = {".txt", ".docx"}


def read_transcript(path: str | Path) -> str:
    source = Path(path)
    ext = source.suffix.lower()
    if ext == ".txt":
        return source.read_text(encoding="utf-8-sig")
    if ext == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX support requires python-docx. Run: pip install -r requirements.txt") from exc
        document = Document(str(source))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    raise ValueError(f"Unsupported input file type: {ext}")


def write_text(path: str | Path, text: str) -> None:
    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(text, encoding="utf-8")


def write_docx(path: str | Path, text: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX output requires python-docx. Run: pip install -r requirements.txt") from exc

    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    document.save(str(destination))


def ensure_directories(paths: Iterable[str | Path]) -> None:
    for path in paths:
        Path(path).mkdir(parents=True, exist_ok=True)
